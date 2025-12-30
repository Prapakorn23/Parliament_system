"""
Document Processor for Word Documents (.doc, .docx)
ระบบแปลงไฟล์ Word Document เป็น text
"""

import os
from typing import Tuple, Optional
import io

# Word Document Processing Libraries
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

try:
    # สำหรับ .doc เก่า (Office 97-2003 format)
    # ใช้ผ่าน textract หรือ win32com
    pass
except ImportError:
    pass


class DocProcessor:
    """ประมวลผลไฟล์ Word Document และแปลงเป็น text"""
    
    def __init__(self):
        self.supported_methods = self._check_available_libraries()
    
    def _check_available_libraries(self) -> dict:
        """ตรวจสอบ libraries ที่สามารถใช้งานได้"""
        return {
            'python-docx': PYTHON_DOCX_AVAILABLE,  # สำหรับ .docx
            'textract': TEXTRACT_AVAILABLE  # สำหรับ .doc และ .docx (ต้องติดตั้ง dependencies เพิ่ม)
        }
    
    def extract_text_from_docx(self, docx_path: str) -> Tuple[bool, str, str]:
        """
        แปลง .docx เป็น text ด้วย python-docx
        
        Args:
            docx_path: path ของไฟล์ .docx
            
        Returns:
            Tuple of (success, text, method_used)
        """
        if not PYTHON_DOCX_AVAILABLE:
            return False, "", "python-docx library not installed"
        
        try:
            doc = Document(docx_path)
            full_text = []
            
            # อ่านข้อความจาก paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # อ่านข้อความจาก tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            text = '\n'.join(full_text)
            return True, text, "python-docx"
            
        except Exception as e:
            return False, "", f"python-docx error: {str(e)}"
    
    def extract_text_from_doc(self, doc_path: str) -> Tuple[bool, str, str]:
        """
        แปลง .doc (เก่า) เป็น text
        ใช้ textract หรือวิธีอื่น
        
        Args:
            doc_path: path ของไฟล์ .doc
            
        Returns:
            Tuple of (success, text, method_used)
        """
        # วิธีที่ 1: ใช้ textract (รองรับหลาย format)
        if TEXTRACT_AVAILABLE:
            try:
                text = textract.process(doc_path).decode('utf-8')
                return True, text, "textract"
            except Exception as e:
                error_msg = str(e)
                # ถ้า textract ไม่สำเร็จ ลองวิธีอื่น
                pass
        
        # วิธีที่ 2: ใช้ python-docx (อาจใช้ไม่ได้กับ .doc เก่า)
        # .doc เก่ามักใช้ไม่ได้กับ python-docx
        
        # วิธีที่ 3: แนะนำให้แปลงเป็น .docx ก่อน
        return False, "", "ไม่สามารถอ่านไฟล์ .doc ได้ กรุณาแปลงเป็น .docx ก่อน หรือติดตั้ง textract"
    
    def extract_text_from_document(self, doc_path: str) -> Tuple[bool, str, str]:
        """
        แปลง Word Document เป็น text โดย auto-detect ประเภทไฟล์
        
        Args:
            doc_path: path ของไฟล์ (.doc หรือ .docx)
            
        Returns:
            Tuple of (success, text, method_used)
        """
        if not os.path.exists(doc_path):
            return False, "", "ไม่พบไฟล์"
        
        filename_lower = doc_path.lower()
        
        # รองรับ .docx
        if filename_lower.endswith('.docx'):
            return self.extract_text_from_docx(doc_path)
        
        # รองรับ .doc (เก่า)
        elif filename_lower.endswith('.doc'):
            return self.extract_text_from_doc(doc_path)
        
        else:
            return False, "", "ไม่ใช่ไฟล์ Word Document (.doc หรือ .docx)"
    
    def get_installation_instructions(self) -> dict:
        """คำแนะนำการติดตั้ง libraries"""
        instructions = {
            'basic': {
                'missing': [],
                'command': []
            },
            'textract': {
                'missing': [],
                'command': [],
                'note': ''
            }
        }
        
        # ตรวจสอบ basic libraries
        if not PYTHON_DOCX_AVAILABLE:
            instructions['basic']['missing'].append('python-docx')
        
        if instructions['basic']['missing']:
            instructions['basic']['command'] = f"pip install {' '.join(instructions['basic']['missing'])}"
        
        # ตรวจสอบ textract (สำหรับ .doc เก่า)
        if not TEXTRACT_AVAILABLE:
            instructions['textract']['missing'].append('textract')
            instructions['textract']['command'] = "pip install textract"
            instructions['textract']['note'] = "สำหรับ .doc เก่า ต้องติดตั้ง dependencies เพิ่มเติม (เช่น antiword สำหรับ Linux)"
        
        return instructions
    
    def check_doc_type(self, doc_path: str) -> str:
        """
        ตรวจสอบว่าเป็นไฟล์ .doc หรือ .docx
        
        Returns:
            'docx', 'doc', หรือ 'unknown'
        """
        filename_lower = doc_path.lower()
        if filename_lower.endswith('.docx'):
            return 'docx'
        elif filename_lower.endswith('.doc'):
            return 'doc'
        else:
            return 'unknown'

